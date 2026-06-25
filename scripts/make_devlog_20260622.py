"""2026-06-22 개발일지를 Word(.docx)로 생성 → 프로젝트 + Windows host 양쪽 저장."""

import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DATE = "2026-06-22"
FILENAME = f"블로그_개발일지_{DATE}.docx"
PROJECT_PATH = f"/home/es0764/blog-platform/{FILENAME}"

doc = Document()

# 한글이 잘 보이도록 기본 폰트를 맑은 고딕으로
normal = doc.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


# ===== 표지 =====
title = doc.add_heading("블로그 플랫폼 개발일지", level=0)
sub = doc.add_paragraph()
r = sub.add_run(f"{DATE} · 스택: FastAPI + PostgreSQL + React + Docker")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p(
    "오늘 주제: 통합 포털 / 서비스 상태·업타임 / AI 글 초안(코드) / "
    "Docker 전체 컨테이너화. 한마디로 '기능 마무리 + 인프라 첫걸음'.",
    bold=True,
)

# ===== 0. 오늘 요약 =====
h1("0. 오늘 요약")
bullet("통합 사이트(포털): 루트(/)에 로비를 두고 블로그(/blog)·상태(/status)로 갈라지는 입구 신설.")
bullet("서비스 상태 + 업타임: 백엔드/DB/메일 점검 + 가동률 %를 서비스별 날짜 막대로 표시.")
bullet("AI 글 초안: 거친 메모 → 정돈된 글 구조 마크다운. 코드 완료, API 키만 넣으면 동작.")
bullet("Docker 전체 컨테이너화(오늘 메인): `docker compose up` 한 줄로 4개 컨테이너 전체 기동.")

# ===== 1. 통합 포털 =====
h1("1. 통합 사이트(포털)")
p("비유: 회사 건물에 들어가면 바로 사무실이 아니라 로비 안내판에서 부서별 층을 보고 가듯, "
  "사이트 루트에 '로비(포털)'를 두고 블로그/상태정보로 안내하는 구조.")
h2("무엇을 했나")
bullet("/ = 통합 랜딩(블로그·상태정보 카드 2개), 기존 블로그 홈은 /blog 로 이동.")
bullet("라우트 재배치: /blog, /blog/posts/:id(+/edit), /blog/new, /status, /login, /register.")
bullet("라우팅을 옮기면 내부 링크가 깨지므로 글 링크·'목록으로'·작성 후 이동 등을 전부 함께 수정.")
h2("실무 포인트")
bullet("URL을 /blog 아래로 모으면 이후 기능 추가(/shop, /docs 등) 시에도 구조가 깔끔하다.")

# ===== 2. 상태 + 업타임 =====
h1("2. 서비스 상태 페이지 + 업타임")
p("비유: 지하철역의 '에스컬레이터 가동중/점검중' 전광판, 또는 claude status / github status 페이지.")
h2("무엇을 했나")
bullet("백엔드 /status: backend·database·mail 점검 + 통계(글 수·구독자 수) 반환.")
bullet("메일 점검은 Mailpit SMTP(1025) 소켓 연결로, DB는 select 1 + count 쿼리로.")
bullet("업타임: status_checks 테이블에 1분마다 백그라운드 스레드가 자가 점검을 1줄씩 기록.")
bullet("/status/history?days=N: 일별로 '정상 점검 ÷ 전체 점검' 비율을 서비스별로 집계.")
bullet("프론트: 백엔드/DB/메일 각각 막대(초록=정상, 노랑=일부장애, 빨강=장애, 회색=데이터없음) + 전체 %.")
h2("실무 포인트 (정직한 한계)")
bullet("자기 자신이 죽으면 자기가 기록을 못 남긴다 → 우리 업타임은 '돌고 있을 때 건강했나' 기준.")
bullet("진짜 다운타임 측정은 '밖에서 찔러보는 외부 감시자'가 필요(inside-out vs outside-in 모니터링).")
bullet("외부 모니터는 이후 AWS 단계(Lambda + 크론)에서 진짜로 붙일 수 있음.")

# ===== 3. AI 초안 =====
h1("3. AI 글 초안 생성 (코드 완료 / 키 대기)")
p("비유: 회의 후 휘갈긴 메모를 비서가 깔끔한 보고서 골격으로 정리해 주는 것.")
h2("무엇을 했나")
bullet("백엔드 POST /ai/draft: 메모를 받아 Claude API로 '제목·소제목·초안' 마크다운 생성.")
bullet("공식 anthropic SDK 사용, 모델은 config.ai_model(기본 claude-opus-4-8, .env로 교체 가능).")
bullet("프론트: 글쓰기 페이지에 'AI 초안' 박스(메모 입력 → 생성 → 첫 # 제목은 제목칸, 나머지는 본문).")
h2("실무 포인트")
bullet("비용 보호: 유료 API라 아무나 호출 못 하게 로그인 필수로 제한(봇이 비용을 태우는 것 방지).")
bullet("시크릿 관리: 키는 .env(gitignore)에만, 코드/커밋 절대 금지. 모델명도 env로 분리.")
bullet("현재 키 미설정 → 호출 시 친절한 503 안내. $5 충전 + 키 입력이면 즉시 동작.")

# ===== 4. Docker =====
h1("4. Docker 전체 컨테이너화 (오늘 메인)")
p("비유: 이삿짐을 규격 컨테이너에 담으면 트럭→배→기차 어디든 그대로 옮겨지듯, "
  "앱을 컨테이너로 싸면 '내 노트북에서 되던 것'이 AWS 서버에서도 그대로 된다.")

h2("핵심 개념")
bullet("이미지 = 냉동만두(설계도/스냅샷), 컨테이너 = 쪄낸 만두(실행 인스턴스). 클래스 vs 인스턴스와 같음.")
bullet("레이어 캐시: Dockerfile 한 줄=필름 한 장. 안 바뀐 아래 층은 재사용 → requirements 설치를 코드 복사보다 위(아래 층)에 둠.")
bullet("네트워킹: 컨테이너 안 localhost는 '자기 자신'. DB/메일은 compose 서비스 이름(db, mailpit)으로 접속. AWS에선 이 자리에 RDS 주소가 들어감.")
bullet("볼륨: named volume(postgres_data)=USB 같은 영속 저장소 → 컨테이너 지워도 데이터 유지. bind mount(./backend:/app)=공유폴더(코드 즉시 반영).")
bullet("healthcheck: DB는 컨테이너가 켜져도 몇 초간 연결을 못 받음 → pg_isready로 '준비 완료' 확인 후 백엔드 기동(경합 방지).")

h2("1단계 — 백엔드 컨테이너")
bullet("backend/.dockerignore: .venv·uploads·.env 제외(경량화 + 시크릿 보호).")
bullet("compose backend: env_file(없던 .env) 제거 → environment로 DATABASE_URL=@db, SMTP_HOST=mailpit 주입.")
bullet("db healthcheck(pg_isready) + depends_on: service_healthy.")
bullet("command: 'alembic upgrade head && uvicorn ...' → 시작 시 마이그레이션 자동.")

h2("2단계 — 프론트 컨테이너 + 전체 한 방")
bullet("멀티스테이지 Dockerfile: node로 빌드(dist 생성) → nginx가 정적 서빙(최종 이미지엔 node/소스 없음).")
bullet("nginx.conf SPA 폴백(try_files … /index.html): /blog 등 새로고침해도 404 안 남.")
bullet("포트 5173:80(nginx 80 → 호스트 5173) → CORS origin(localhost:5173) 유지, 백엔드 안 고침.")
bullet("AWS S3+CloudFront 정적 배포와 동일한 구조 → 이후 이전이 수월.")

# ===== 5. 검증 결과 =====
h1("5. 검증 결과 (자가 검증)")
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "항목"
table.rows[0].cells[1].text = "결과"
checks = [
    ("프론트 빌드/lint", "전부 통과(에러 0)"),
    ("라우트 200", "/ /blog /status /blog/posts/1 전부 200(SPA 폴백 동작)"),
    ("백엔드 /status", "backend·database·mail = ok, 통계 posts:4/subscribers:1"),
    ("업타임 기록", "status_checks 1분마다 누적, 서비스별 집계 정상"),
    ("AI /ai/draft", "비로그인 401 / 로그인+키없음 503(친절 안내)"),
    ("Docker 스택", "compose up 한 줄 → 4컨테이너(db healthy/mailpit/backend/frontend) Up"),
    ("데이터 보존", "컨테이너 재생성에도 posts:4 유지(named volume)"),
    ("CORS", "access-control-allow-origin: http://localhost:5173 확인"),
]
for a, b in checks:
    row = table.add_row().cells
    row[0].text = a
    row[1].text = b

# ===== 6. 자주 쓰는 명령 =====
h1("6. 자주 쓰는 Docker 명령")
for c in [
    "docker compose up -d --build   # 전체 빌드+기동(백그라운드)",
    "docker compose ps              # 상태 확인",
    "docker compose logs -f backend # 로그 보기",
    "docker compose down            # 정지(데이터 유지)",
    "docker compose down -v         # 데이터까지 삭제(주의)",
    "docker compose up -d --build frontend  # 프론트 코드 수정 후 다시 굽기",
]:
    bullet(c)
p("주의: 현재 셸은 docker 그룹 미반영 → 실제로는 sg docker -c \"...\"로 감싸 실행. "
  "백엔드 포트 종료는 fuser -k 8000/tcp (pkill은 자기 셸까지 죽임).")

# ===== 7. 다음 할 일 =====
h1("7. 다음 할 일")
bullet("(선택) AI 키: 루트 .env에 ANTHROPIC_API_KEY=sk-... 넣고 백엔드 재기동 → 초안 생성 e2e.")
bullet("AWS 수동 배포: 프론트 S3+CloudFront / 백엔드 컨테이너+RDS / 도메인·HTTPS. (비용 발생 — 예상비용·프리티어부터 설명)")
bullet("이후: Terraform(인프라 코드화) → GitHub Actions CI/CD(push→자동 배포).")
bullet("(별도) 7단계: Lambda + Bedrock 사이트 내장 AI. (Bedrock은 AWS 결제 — Anthropic 키와 무관)")

# ===== 저장 =====
doc.save(PROJECT_PATH)
print(f"[저장] 프로젝트: {PROJECT_PATH}")

# Windows host 저장 위치 선택 (Desktop 우선, 없으면 Documents, 없으면 사용자 폴더 루트)
host_user = "/mnt/c/Users/erert"
candidates = [
    os.path.join(host_user, "Desktop"),
    os.path.join(host_user, "OneDrive", "Desktop"),
    os.path.join(host_user, "바탕 화면"),
    os.path.join(host_user, "Documents"),
    host_user,
]
host_dir = next((d for d in candidates if os.path.isdir(d)), host_user)
host_path = os.path.join(host_dir, FILENAME)
shutil.copy2(PROJECT_PATH, host_path)
print(f"[저장] Windows host: {host_path}")
