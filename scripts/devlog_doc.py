"""개발일지 docx의 **공통 껍데기** — 서식·헬퍼·저장/복사.

## 왜 이걸 만들었나

`make_devlog_YYYYMMDD.py`가 29개 있고 합치면 8,775줄인데, 그중 파일마다 반복되는
서식 코드가 130줄쯤이었다(스타일 지정 · h1/h2/p/bullet/analogy/field/note/warn/ev ·
저장과 바탕화면 복사). 즉 **본문보다 껍데기를 더 많이 복사해 왔다.**

2026-08-11 동료 리뷰가 이걸 지적했을 때 바로 안 고친 이유가 있다: 기존 28편을
새 코드로 **재생성하면 안 되기 때문**이다. 이미 발행된 글이고, 서식이 한 픽셀이라도
달라지면 "그때 그 문서"가 아니게 된다. 그래서 결론이 **"기존은 손대지 말고 새 편부터"**
였고, #29가 그 첫 편이다. 이 파일을 만들면서 옛 파일은 하나도 건드리지 않았다.

## 쓰는 법

    from devlog_doc import DevlogDoc

    d = DevlogDoc("2026-08-11", "블로그 만들기 #29")
    d.h1("1. 아침")
    d.p("...")
    d.ev("실제 출력 — ...", ["...", "..."])
    d.save()

실행 (WSL 파이썬에 python-docx가 없어서 컨테이너로 만든다):

    docker run --rm -v "$(wslpath -w "$PWD")":/w \\
      -v "/mnt/c/Users/USER/Desktop/개발일지":/desktop \\
      -w /w python:3.12 sh -c "pip install -q python-docx && python scripts/make_devlog_YYYYMMDD.py"

## 서식을 바꾸려면

**여기를 바꾸면 앞으로 쓸 편만 바뀐다** — 그게 의도다. 이미 나간 편의 서식은
그 편의 스크립트가 아니라 **이미 배포된 docx와 content/devlog/*.md**가 진실이다.
"""

import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7E, 0x34)
MONO = RGBColor(0x24, 0x29, 0x2E)

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# 바탕화면의 개발일지 폴더. 컨테이너 안에서는 /desktop으로 마운트된다.
# **바탕화면 루트에 흩뿌리지 않는다** — 폴더 안에만 둔다.
# 바탕화면 사본을 둘 자리. **순서가 중요하다** — 첫 번째로 존재하는 곳 하나에만 복사한다.
#
# ⚠️ 이 컴퓨터의 바탕화면은 **OneDrive로 리디렉션돼 있다.** 사용자가 실제로 보는 폴더는
# `OneDrive/바탕 화면/개발일지`이고, `C:\Users\USER\Desktop\개발일지`는 **다른 폴더**다.
# 그런데 목록의 첫 자리가 후자였던 탓에 #32·#33이 사용자 눈에 안 보이는 곳으로 갔다
# (2026-08-17에 "개발일지가 14일까지밖에 없다"고 해서 알았다 — 파일은 멀쩡히 만들어졌고
# 복사 로그도 찍혔다. 만든 것과 닿은 것이 다른, 그날 하루 종일 나온 그 모양이다).
# OneDrive 쪽을 앞에 둔다. `/desktop`은 컨테이너로 만들 때 마운트하는 자리라 최우선.
HOST_DIRS = (
    "/desktop",
    "/mnt/c/Users/USER/OneDrive/바탕 화면/개발일지",
    "/mnt/c/Users/USER/Desktop/개발일지",
)


class DevlogDoc:
    """개발일지 한 편. 서식은 전부 여기가 쥐고, 호출부는 내용만 쓴다."""

    def __init__(self, date: str, series: str):
        self.date = date
        self.series = series
        self.filename = f"블로그_개발일지_{date}.docx"
        self.path = os.path.join(REPO_ROOT, self.filename)

        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = "맑은 고딕"
        normal.font.size = Pt(10.5)
        # 한글은 eastAsia 글꼴을 따로 지정하지 않으면 Word가 제멋대로 고른다.
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    @staticmethod
    def _strip_dup(label: str, text: str) -> str:
        """호출부가 접두사를 한 번 더 넣었으면 떼어낸다.

        왜 필요한가 — 2026-08-14에 #30 본문 첫 문장이 "이번 편의 형식: 이번 편의 형식:"
        으로 라이브에 나가 있는 걸 발견했다. 캡션도 "▶ 실제 출력 — 실제 출력 —"이었다.
        이 클래스는 #29부터 쓰기 시작한 공통 서식인데, **접두사를 여기서 붙인다는 걸
        호출부가 모르고 같은 말을 또 썼다.** 그 편의 발췌문이 홈·아카이브·RSS 세 곳의
        맨 위에 뜨는 자리라 가장 눈에 띄는 오타가 됐다.

        막는 대신 **떼어내고 알린다**: 생성기는 사람이 손으로 돌리는 도구라 예외로
        멈추면 그 자리에서 고치느라 흐름이 끊기고, 조용히 두면 이번처럼 발행까지 간다.
        """
        stripped = text.lstrip()
        if stripped.startswith(label):
            print(f"  ⚠️ 접두사 중복 제거: '{label}' — 호출부에서 그 말을 빼세요.")
            return stripped[len(label) :].lstrip()
        return text

    # ── 표지 ────────────────────────────────────────────────────────────
    def cover(self, subject: str, audience: str, note: str | None = None) -> None:
        self.doc.add_heading("블로그 개발일지", level=0)
        self.p(f"날짜: {self.date}   ·   {self.series}", bold=True)
        self.p(f"주제: {subject}")
        self.p(f"대상: {audience}")
        if note:
            self.p(f"이번 편의 형식: {self._strip_dup('이번 편의 형식:', note)}")

    # ── 문단 ────────────────────────────────────────────────────────────
    def h1(self, t: str) -> None:
        self.doc.add_heading(t, level=1)

    def h2(self, t: str) -> None:
        self.doc.add_heading(t, level=2)

    def p(self, t: str, bold: bool = False):
        para = self.doc.add_paragraph()
        para.add_run(t).bold = bold
        return para

    def bullet(self, t: str) -> None:
        self.doc.add_paragraph(t, style="List Bullet")

    # ── 색이 붙는 상자들 ────────────────────────────────────────────────
    def _tagged(self, tag: str, text: str, color) -> None:
        para = self.doc.add_paragraph()
        r = para.add_run(tag)
        r.bold = True
        r.font.color.rgb = color
        r2 = para.add_run(text)
        r2.font.color.rgb = color
        para.paragraph_format.left_indent = Pt(12)

    def analogy(self, t: str) -> None:
        self._tagged("🔎 비유   ", t, BLUE)

    def note(self, t: str) -> None:
        self._tagged("🛠 전문가 노트   ", t, GREEN)

    def warn(self, t: str) -> None:
        self._tagged("⚠️ 함정   ", t, RED)

    def field(self, label: str, text: str, color=GRAY):
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(12)
        r = para.add_run(f"{label}  ")
        r.bold = True
        r.font.color.rgb = color
        para.add_run(text)
        return para

    # ── 터미널 증거 ─────────────────────────────────────────────────────
    def ev(self, title: str, lines: list[str]) -> None:
        """실제로 화면에 찍힌 문자열. **지어내지 않는다** — 이 시리즈의 규칙이다."""
        cap = self.doc.add_paragraph()
        cap.paragraph_format.left_indent = Pt(12)
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(f"▶ 실제 출력 — {self._strip_dup('실제 출력 —', title)}")
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = GRAY

        for line in lines:
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(line if line else " ")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = MONO
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        self.doc.add_paragraph("").paragraph_format.space_after = Pt(2)

    # ── 저장 ────────────────────────────────────────────────────────────
    def save(self) -> None:
        self.doc.save(self.path)
        print(f"저장: {self.path}")
        for host_dir in HOST_DIRS:
            if os.path.isdir(host_dir):
                shutil.copy(self.path, os.path.join(host_dir, self.filename))
                print(f"복사: {host_dir}/{self.filename}")
                break
        else:
            print("호스트 개발일지 폴더를 못 찾았다 — 저장소 사본만 만들었다.")
