"""2026-06-24 개발일지(AWS 풀스택 배포)를 Word로 → 프로젝트 + Windows host 양쪽 저장."""

import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DATE = "2026-06-24"
FILENAME = f"블로그_개발일지_{DATE}.docx"
PROJECT_PATH = f"/home/es0764/blog-platform/{FILENAME}"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def p(t, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.bold = bold
    return para


def quote(t):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


# ===== 표지 =====
doc.add_heading("블로그 플랫폼 개발일지", level=0)
sub = doc.add_paragraph()
r = sub.add_run(f"{DATE} · 오늘의 주제: AWS 풀스택 배포 (대장정)")
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
p("로컬에서 docker compose로 돌던 블로그를, 진짜 인터넷(AWS)에 HTTPS로 올린 날. "
  "AWS CLI 설치부터 RDS·EC2·Docker·CloudFront 라우팅, 마지막 WAF 디버깅까지.", bold=True)
p("라이브 URL: https://d2j66m9udyg9yq.cloudfront.net", bold=True)

# ===== 0. 요약 =====
h1("0. 오늘 한 일 요약")
bullet("AWS CLI 설치 + IAM 키 등록 (내 컴퓨터 ↔ AWS 연결)")
bullet("RDS(PostgreSQL) 생성 — 데이터 창고")
bullet("EC2(t2.micro) 생성 + Docker 설치 — 백엔드 돌릴 컴퓨터")
bullet("백엔드 컨테이너 배포 + RDS 연결 + 마이그레이션")
bullet("CloudFront가 /api를 EC2로 라우팅 → 단일 HTTPS로 통합 (혼합콘텐츠·CORS 해결)")
bullet("프론트 재배포(S3 sync + 캐시 무효화) → 전체 연결")
bullet("WAF가 이미지 업로드 막던 문제 디버깅 & 해결")

# ===== 1. 큰 그림 (식당 비유) =====
h1("1. 큰 그림 — '식당 하나' 비유")
p("우리 블로그를 식당 하나로 보면 이해가 쉽다.")
bullet("화면(프론트, S3+CloudFront) = 간판·메뉴판·인테리어 (손님 눈에 보이는 것)")
bullet("백엔드(EC2) = 주방 (주문 받아 요리 = 데이터 처리)")
bullet("DB(RDS) = 식자재 창고 (글·계정 보관)")
bullet("보안 그룹 = 출입문 문지기 (누가 어느 문으로 들어오나)")
quote("핵심: 로컬에서 docker compose로 띄우던 '주방+창고'를, 클라우드의 EC2(주방)와 "
      "RDS(전문 창고)로 옮긴 것. 코드는 그대로, 장소만 내 노트북 → AWS로 바뀜.")
p("그림으로:")
p("브라우저 ──HTTPS──► CloudFront(안내데스크) ─┬─ 화면 → S3(창고)\n"
  "                                              └─ /api → EC2(주방) → RDS(식자재 창고)")

# ===== 2. AWS CLI =====
h1("2. AWS CLI — '리모컨'")
p("비유: 콘솔 클릭 = 자판기 앞에서 매번 손으로 누르기. CLI = 명령 한 줄로 같은 걸 즉시.")
bullet("브라우저 클릭 대신 터미널 명령으로 AWS를 다루는 도구")
bullet("예: S3 업로드를 'aws s3 sync dist/ s3://버킷' 한 줄로 (드래그·폴더 싸움 끝)")
bullet("진짜 이유: 클릭은 재현이 안 됨. 명령은 저장·반복·자동화 가능 → CI/CD·Terraform의 기반")
bullet("연결엔 IAM 사용자 + 액세스 키 필요 (시크릿이라 ~/.aws에만, 깃 금지)")

# ===== 3. RDS =====
h1("3. RDS — 'DB 관리인이 딸린 창고'")
p("비유: Postgres를 직접 설치·백업하는 대신, AWS가 다 관리해주는 DB를 빌림.")
bullet("PostgreSQL, db.t3.micro, 단일 AZ, 20GB (전부 프리티어 무료)")
bullet("퍼블릭 액세스 = 아니요(비공개) → 인터넷에 직접 노출 X, EC2만 접근")
bullet("왜 EC2 안에 안 두고 따로? → 주방에 불나도 창고(데이터)는 멀쩡하게. 안전·정석.")
bullet("실수: '초기 DB 이름 blog'를 빠뜨림 → prod는 기본 postgres DB를 사용하는 걸로 우회(테이블은 alembic이 생성)")

# ===== 4. EC2 =====
h1("4. EC2 — '클라우드에 빌린 컴퓨터(주방)'")
bullet("Amazon Linux 2023, t2.micro (프리티어 750h = 한 달 내내 1대 무료)")
bullet("키 페어(blog-key.pem) = 주방 들어가는 열쇠 (SSH 접속용)")
bullet("보안 그룹 규칙: SSH 22(내 IP만) + API 8000(누구나)")
bullet("자동 할당 퍼블릭 IP로 인터넷에서 접근 (NAT 게이트웨이 안 씀 = 비용 함정 회피)")

# ===== 5. 백엔드 배포 =====
h1("5. 백엔드 배포 — '이삿짐을 규격 컨테이너에'")
p("비유: 이삿짐을 규격 컨테이너에 담으면 트럭→배→기차 어디든 그대로 옮겨진다. "
  "도커도 똑같아서, 내 노트북에서 되던 게 AWS에서도 '그대로' 된다 ('제 PC에선 됐는데요' 박멸).")
bullet("핵심 깨달음: 컨테이너 안 localhost는 '자기 자신' → DB는 서비스명/RDS 엔드포인트로 접속")
bullet("EC2에 Docker 설치 → 백엔드 코드 전송(tar+scp) → docker compose로 빌드·실행")
bullet("시작 시 alembic 마이그레이션 자동 → RDS에 테이블 7개 생성")
bullet("RDS 비번은 EC2의 .env에만 (사용자가 직접 입력, AI는 안 봄)")

# ===== 6. CloudFront 라우팅 =====
h1("6. CloudFront /api 라우팅 — '안내 데스크' (오늘의 핵심 마법)")
p("문제: 프론트는 https(CloudFront), 백엔드는 http(EC2:8000). 브라우저는 'HTTPS 페이지가 "
  "HTTP 부르는 것'을 차단(혼합 콘텐츠) + 다른 도메인이라 CORS도 필요.")
p("해결: CloudFront를 모든 요청의 '안내 데스크'로 세우고, 경로(path)로 분배.")
bullet("/, /blog, /status … (화면) → S3 (정적 파일)")
bullet("/api/* (데이터) → EC2 (백엔드)")
bullet("/uploads/* (이미지) → EC2")
quote("브라우저 입장에선 처음부터 끝까지 같은 HTTPS 주소 한 곳과만 대화 → 혼합콘텐츠 없음, "
      "CORS 없음. 실제 http 통신은 CloudFront↔EC2 내부에서만.")
bullet("그래서 백엔드 라우트를 전부 /api 밑으로 옮김(main.py prefix='/api')")
bullet("/api/* = 캐시 끔(매번 최신) + AllViewer(로그인 토큰 전달), /uploads/* = 캐시 켬(이미지)")
bullet("이 'behavior(경로별 규칙)'를 직접 콘솔에서 삭제→재생성하며 익힘")

# ===== 7. WAF 디버깅 =====
h1("7. WAF 디버깅 — '문지기가 큰 짐을 막은 사건' (실무급 트러블슈팅)")
p("증상: 큰 이미지(8KB+) 업로드 시 'Unexpected token < ... not valid JSON' 에러. "
  "작은 건 됨. EC2 직접은 성공인데 CloudFront 경유만 실패.")
p("추적 과정 (이게 진짜 실무 디버깅):")
bullet("① EC2 직접 200, CloudFront는 HTML(index.html) → 어딘가 에러를 폴백으로 바꿔치기")
bullet("② 응답 헤더 server:AmazonS3 + x-cache:Error → 403 폴백(우리가 만든 403→index.html)")
bullet("③ 크기로 좁힘: 60바이트 OK, 13KB 막힘 → '본문 크기 차단' = WAF 전형 증상")
bullet("④ 배포에 WAF(CreatedByCloudFront) 연결 확인 → CommonRuleSet의 SizeRestrictions_BODY(8KB 초과 차단)가 범인")
p("해결:")
bullet("CloudFront '무료 보안' = 사실 WAF 켜진 것이었음 (무료=끄기가 아니라 무료등급 WAF 켜기)")
bullet("플랜 구독 중이라 WAF 통째 제거는 막힘 → 그 룰만 '차단→Count(통과)'로 튜닝(update-web-acl)")
bullet("결과: 13KB·1MB 업로드 200, 이미지 표시까지 정상")
quote("교훈: 프로덕션에서 WAF가 조용히 기능을 막는 건 엔지니어가 자주 겪는 함정. "
      "'로컬은 되는데 배포만 안 됨 → 크기로 좁히기 → WAF 발견 → 룰 튜닝' 전 과정을 경험.")

# ===== 8. 비용 =====
h1("8. 비용")
bullet("EC2·RDS·S3·CloudFront 전부 프리티어 → 12개월간 사실상 $0")
bullet("주의: WAF는 '무료' 플랜이지만 확실치 않음 → Billing에서 며칠 뒤 확인 권장")
bullet("예산 알람 설정해둠($10). 추천: $1짜리도 추가하면 작은 누수(WAF 등)도 즉시 포착")
bullet("나중 메일/도메인: 도메인 연 ~$12, SES 발송 거의 $0")
bullet("안 쓸 땐 EC2·RDS stop으로 절약. 단 stop하면 IP 바뀌어 CloudFront 오리진 깨짐(EIP 미사용)")

# ===== 9. 검증 =====
h1("9. 검증 결과")
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "항목"
table.rows[0].cells[1].text = "결과"
for a, b in [
    ("프론트 화면 /, /blog", "CloudFront 200"),
    ("API /api/status", "database:ok (RDS 연결)"),
    ("API /api/posts", "200 (빈 목록 [])"),
    ("로그인/인증", "AllViewer로 토큰 전달 → 동작"),
    ("이미지 업로드(WAF 튜닝 후)", "13KB·1MB 200, 표시 OK"),
    ("메일", "down (prod 메일서버 없음 — 의도된 미구현, SES 나중)"),
]:
    c = table.add_row().cells
    c[0].text = a
    c[1].text = b

# ===== 10. 배운 것 =====
h1("10. 오늘 배운 핵심 (면접 답변용)")
quote('"정적 프론트는 S3+CloudFront(OAC)로, 동적 백엔드는 EC2에 Docker로, DB는 RDS를 '
      '비공개로 분리 배포했습니다. CloudFront에서 /api를 백엔드 오리진으로 라우팅해 단일 '
      'HTTPS 도메인으로 통합, 혼합 콘텐츠와 CORS를 동시에 해결했습니다. 배포 후 WAF가 큰 '
      '파일 업로드를 막는 문제를 진단해 룰을 튜닝했습니다."')
bullet("S3+CloudFront(OAC) 정적 배포 — 직접 콘솔로")
bullet("EC2 + RDS 분리 배포, 보안 그룹(문지기) 설계")
bullet("CloudFront behavior(경로 라우팅) — 직접 삭제·재생성")
bullet("실무급 디버깅(WAF) — 증상→격리→원인→해결 전 과정")

# ===== 11. 다음 =====
h1("11. 다음 할 일")
bullet("Terraform — 오늘 손으로 만든 인프라를 코드로 (만들고/부수기 자유). 다음 차례.")
bullet("GitHub Actions CI/CD — push하면 자동 배포")
bullet("(선택) 도메인 구매 + SES 메일 + EIP(IP 고정)")
bullet("(대기) AI 글 초안 — ANTHROPIC_API_KEY 넣으면 켜짐")

# ===== 저장 =====
doc.save(PROJECT_PATH)
print(f"[저장] 프로젝트(WSL): {PROJECT_PATH}")

host_user = "/mnt/c/Users/erert"
cands = [
    os.path.join(host_user, "Desktop"),
    os.path.join(host_user, "OneDrive", "Desktop"),
    os.path.join(host_user, "바탕 화면"),
    os.path.join(host_user, "Documents"),
    host_user,
]
host_dir = next((d for d in cands if os.path.isdir(d)), host_user)
host_path = os.path.join(host_dir, FILENAME)
shutil.copy2(PROJECT_PATH, host_path)
print(f"[저장] Windows host: {host_path}")
