"""마크다운으로 쓴 개발일지를 docx로 옮긴다.

## 왜 필요한가

#29부터 #35까지는 `make_devlog_YYYYMMDD.py`가 **본문을 파이썬으로 쓰고** docx를 만든
다음, `devlog_to_markdown.py`가 그 docx에서 마크다운을 뽑는 순서였다. 그런데 08-26부터
글을 마크다운으로 바로 쓰기 시작하면서 그 순서가 뒤집혔고, 생성기 스크립트가 없는 편이
셋 쌓였다(08-26·08-27·08-28). 그래서 저장소에는 글이 있는데 **바탕화면 폴더에는 그 편이
없는** 상태가 됐다.

이 스크립트가 그 방향을 메운다. 서식은 `devlog_doc.DevlogDoc`이 그대로 쥐고 있으므로
기존 편들과 같은 모양으로 나온다.

## 무엇을 어떻게 옮기나

    # 블로그 만들기 #38 — 제목        →  날짜 · 시리즈 · 주제(대시 뒤)
    > 입문자가 읽어도 …               →  대상
    이번 편의 형식: …                 →  이번 편의 형식
    ## 1. 소제목                      →  Heading 1
    본문 문단                          →  본문(**굵게**와 `코드`를 살린다)
    - 글머리표                         →  글머리표

**터미널 출력 상자(ev)는 만들지 않는다.** 마크다운에 코드 블록이 있으면 그건 사람이
의도해서 넣은 것이므로 그때 손으로 옮기는 게 낫고, 08-26 이후 세 편에는 코드 블록이
하나도 없다(확인함). 조용히 다른 서식으로 옮기느니 없는 것을 없다고 두는 쪽이다.

사용:
    python scripts/devlog_md_to_doc.py 2026-08-28 [2026-08-27 …]
    python scripts/devlog_md_to_doc.py --missing      # docx가 없는 편 전부
"""

import re
import sys
from pathlib import Path

from devlog_doc import MONO, DevlogDoc
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC_DIR = ROOT / "content" / "devlog"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(para, text: str) -> None:
    """**굵게**와 `코드`를 살려서 문단에 붙인다.

    옛 편들의 docx에는 `**안에**` 처럼 별표가 글자 그대로 남아 있는 자리가 있는데,
    그건 본문을 파이썬 문자열로 쓰면서 마크다운 표기를 그대로 넣었기 때문이다.
    여기서는 서식으로 옮긴다.
    """
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            para.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = para.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = MONO
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            para.add_run(piece)


def convert(date: str) -> Path:
    src = SRC_DIR / f"{date}.md"
    if not src.exists():
        raise SystemExit(f"마크다운이 없다: {src}")

    lines = src.read_text(encoding="utf-8").split("\n")
    if not lines or not lines[0].startswith("# "):
        raise SystemExit(f"첫 줄이 제목(# )이 아니다: {src}")

    title = lines[0][2:].strip()  # "블로그 만들기 #38 — 지웠는데 …"
    series, _, subject = title.partition(" — ")
    audience = ""
    note = ""

    doc = DevlogDoc(date, series.strip())

    body: list[str] = []
    for line in lines[1:]:
        stripped = line.strip()
        if not audience and stripped.startswith("> "):
            audience = stripped[2:]
            continue
        if not note and stripped.startswith("이번 편의 형식:"):
            note = stripped[len("이번 편의 형식:") :].strip()
            continue
        body.append(line)

    doc.cover(subject.strip(), audience, note or None)

    for line in body:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            doc.h1(stripped[3:])
        elif stripped.startswith("### "):
            doc.h2(stripped[4:])
        elif stripped.startswith("- "):
            para = doc.doc.add_paragraph(style="List Bullet")
            add_runs(para, stripped[2:])
        elif stripped.startswith("```"):
            # 코드 블록은 이 경로로 안 옮긴다(머리말 참고). 만나면 알리고 멈춘다.
            raise SystemExit(
                f"{src}에 코드 블록이 있다. ev()로 손으로 옮기고 나서 다시 돌려라."
            )
        else:
            para = doc.doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            add_runs(para, stripped)

    doc.save()
    return Path(doc.path)


def missing_dates() -> list[str]:
    out = []
    for md in sorted(SRC_DIR.glob("20*.md")):
        date = md.stem
        if not (ROOT / f"블로그_개발일지_{date}.docx").exists():
            out.append(date)
    return out


if __name__ == "__main__":
    args = sys.argv[1:]
    if not args:
        raise SystemExit(__doc__.rsplit("사용:", 1)[-1].strip())
    dates = missing_dates() if args == ["--missing"] else args
    if not dates:
        print("docx가 없는 편이 없다.")
    for d in dates:
        convert(d)
