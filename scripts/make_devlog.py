"""오늘(2026-06-21) 개발 내용을 Word(.docx) 문서로 생성."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATE = "2026-06-21"
OUT = f"/home/es0764/blog-platform/블로그_개발일지_{DATE}.docx"

doc = Document()

# 기본 글꼴
doc.styles["Normal"].font.name = "맑은 고딕"
doc.styles["Normal"].font.size = Pt(10.5)


def h1(t):
    doc.add_heading(t, level=1)

def h2(t):
    doc.add_heading(t, level=2)

def p(t):
    doc.add_paragraph(t)

def b(t):
    doc.add_paragraph(t, style="List Bullet")


# ── 표지 ──
title = doc.add_heading("블로그 플랫폼 개발 일지", level=0)
sub = doc.add_paragraph(f"작성일: {DATE}")
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
doc.add_paragraph(
    "스택: FastAPI + PostgreSQL + React(Vite, TypeScript) · 로컬 개발(Docker), 추후 AWS 이전 예정"
)
doc.add_paragraph(
    "목표: 개인 기술 블로그를 직접 만들고 인프라까지 구축하는 학습·포트폴리오 프로젝트."
)

# ── 개요 ──
h1("1. 오늘의 요약")
p("하루 동안 빈 폴더 상태에서 시작해, 풀스택 블로그의 핵심 기능과 디자인을 완성했다. "
  "백엔드(FastAPI)·프론트(React)·DB(PostgreSQL)를 연결하고, 글 CRUD·댓글·이미지·구독·알림·"
  "로그인(JWT)·공개범위·다크모드, 그리고 '글쓴이 구독 시 비공개 글 열람'까지 구현했다.")
b("환경 셋업 → Docker/Postgres → 글 CRUD → 상세+라우팅 → Alembic")
b("이메일 구독 + 새 글 알림(Mailpit)")
b("요약 · 댓글 · 이미지 업로드 · 계정/공개범위(JWT)")
b("로그인/회원가입/글쓰기 페이지 분리")
b("디자인 고급화(Tailwind · Pretendard · 다크모드)")
b("글쓴이 구독 → 일부공개 글 열람")

# ── 단계별 상세 ──
h1("2. 단계별 상세")

h2("2.1 개발 환경 셋업")
p("Python에 pip가 없어 가상환경(venv)을 만들어 FastAPI 의존성을 설치했다. 가상환경을 쓰는 "
  "이유는 프로젝트별로 패키지 버전을 격리해 충돌을 막기 위해서다. 프론트는 Vite로 React+"
  "TypeScript를 초기화했고, 백엔드가 응답하는지(/health)와 브라우저 교차출처 요청을 허용하는지"
  "(CORS) 실제로 호출해 검증했다.")

h2("2.2 Docker / PostgreSQL")
p("Docker Desktop의 WSL 통합이 꺼져 있어 docker 명령이 잡히지 않던 문제를 해결했다. "
  "docker-compose로 postgres:16 컨테이너를 띄워, 내 PC에 DB를 직접 설치하지 않고도 격리된 "
  "데이터베이스를 사용할 수 있게 했다.")

h2("2.3 글 CRUD + 상세 페이지/라우팅")
p("SQLAlchemy 모델로 posts 테이블을 정의하고, 작성/목록/단건/수정/삭제 API를 만들었다. "
  "프론트는 react-router를 도입해 목록(/)과 상세(/posts/:id)를 URL로 분리했다. 글마다 고유 "
  "주소가 생겨 링크 공유가 가능해졌다.")

h2("2.4 Alembic (DB 스키마 버전 관리)")
p("create_all로 즉석 생성하던 테이블을 Alembic 마이그레이션으로 전환했다. Alembic은 'DB 스키마의 "
  "Git'으로, 컬럼 추가·변경을 버전 파일로 추적하고 되돌릴 수 있게 한다. 이후 모든 스키마 변경은 "
  "이 흐름(revision → upgrade)으로 처리했다.")

h2("2.5 이메일 구독 + 새 글 알림(Mailpit)")
p("구독자 이메일을 받는 기능과, 새 글 작성 시 구독자에게 알림 메일을 보내는 기능을 만들었다. "
  "로컬에서는 Mailpit(메일 캐처)으로 실제 발송 없이 보낸 메일을 웹 UI에서 확인했다. 표준 SMTP로 "
  "보내기 때문에 나중에 host/port만 AWS SES로 바꾸면 실제 발송으로 전환된다. 발송은 "
  "BackgroundTasks로 응답 후 비동기 처리해 작성 응답이 지연되지 않도록 했다.")

h2("2.6 요약 · 댓글 · 이미지 업로드")
b("요약: 목록에서 본문 앞부분을 잘라 자동 발췌(DB 변경 없음).")
b("댓글: comments 테이블(글에 외래키, 글 삭제 시 CASCADE로 함께 삭제), 상세 페이지 댓글 UI.")
b("이미지: 업로드 엔드포인트로 파일을 서버에 저장하고 URL을 반환, 본문에 마크다운(![](url))으로 "
  "삽입, 상세 페이지에서 react-markdown으로 렌더링. 저장 위치는 추후 S3로 교체 예정.")

h2("2.7 계정 + 공개범위 (JWT 인증)")
p("회원가입/로그인/내 정보 API를 만들었다. 비밀번호는 bcrypt로 해싱해 평문 저장을 절대 하지 않고, "
  "로그인 시 서명된 JWT 토큰을 발급한다. 토큰은 매 요청의 Authorization 헤더로 신분을 증명하며, "
  "서버가 세션을 따로 보관하지 않아 추후 서버 확장에 유리하다.")
p("글에 작성자(owner)와 공개범위(public/private)를 추가했다. 권한 규칙은 다음과 같다: 작성=로그인 "
  "필수, 비공개 글은 작성자만 조회(없는 것처럼 404로 숨김), 수정·삭제는 소유자만(403). 권한 검사는 "
  "항상 백엔드에서 수행하며, 화면 숨김은 보조 수단이다.")

h2("2.8 로그인/회원가입/글쓰기 페이지 분리")
p("홈에 끼워두던 로그인·글쓰기 폼을 각각 /login, /register, /new(및 /posts/:id/edit) 전용 페이지로 "
  "분리하고, 홈에는 진입 버튼만 두었다. 작업 후에는 자동으로 홈으로 이동한다.")

h2("2.9 디자인 고급화")
b("Tailwind CSS 도입(유틸리티 클래스로 빠른 스타일링) + 타이포그래피 플러그인(prose)으로 본문 렌더링.")
b("Pretendard 웹폰트, indigo 포인트 컬러, 카드형 레이아웃, 부드러운 그림자·hover 효과.")
b("공통 Layout(sticky 헤더/푸터)과 ui.ts 공통 스타일 토큰으로 전 페이지 통일.")
b("다크모드: .dark 클래스 토글 + localStorage 저장 + OS 설정 인식, prose-invert로 본문 반전.")
b("흐릿한 회색 글씨의 대비를 라이트/다크 각각에 맞게 보강해 가독성 향상.")

h2("2.10 글쓴이 구독 → 일부공개 글 열람 (오늘의 마지막 기능)")
p("로그인한 사용자가 특정 글쓴이(작성자)를 구독하면, 그 글쓴이의 일부공개(비공개) 글까지 볼 수 있게 "
  "했다. 이메일 구독과는 다른 '사용자-사용자' 관계라 별도 테이블(author_subscriptions)을 두었다.")
p("핵심 변경은 비공개 글 열람 권한을 '작성자 본인'에서 '작성자 본인 또는 그 작성자를 구독한 사람'으로 "
  "넓힌 것이다. 목록도 DB 쿼리에서 '공개 OR 내 글 OR 내가 구독한 작성자의 글'로 필터링한다. 검증 "
  "결과 같은 비공개 글이 구독 전 404 → 구독 후 200 → 해제 후 다시 404로, 권한이 구독 상태에 정확히 "
  "연동됨을 확인했다.")

# ── 배운 개념 ──
h1("3. 오늘 배운 핵심 개념")
b("Alembic = DB 스키마의 버전 관리(변경 이력을 코드로 추적·복원).")
b("JWT = 서명된 토큰으로 신분 확인. 서버가 세션을 보관하지 않아 확장에 유리.")
b("권한은 항상 백엔드에서 검사. 비공개는 404로 존재 자체를 숨겨 정보 노출을 막음.")
b("CORS = 다른 출처(:5173↔:8000) 간 요청을 서버가 허용하는 헤더 메커니즘.")
b("Tailwind 공통 토큰화 덕에 다크모드를 한 곳 수정으로 전 페이지에 적용.")

# ── DB 구조 ──
h1("4. 현재 데이터베이스 구조")
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "테이블"
hdr[1].text = "설명"
rows = [
    ("users", "회원(이메일 + bcrypt 해시 비밀번호)"),
    ("posts", "글(제목/본문 + owner_id 작성자 + visibility 공개범위)"),
    ("comments", "댓글(글에 외래키, 글 삭제 시 함께 삭제)"),
    ("subscribers", "이메일 구독자(새 글 알림용)"),
    ("author_subscriptions", "글쓴이 구독 관계(구독자→작성자)"),
    ("alembic_version", "현재 적용된 마이그레이션 버전 기록"),
]
for name, desc in rows:
    c = table.add_row().cells
    c[0].text = name
    c[1].text = desc

# ── 다음 할 일 ──
h1("5. 다음 할 일")
b("AI 글 구조 생성(초안 텍스트 → 마크다운 구조).")
b("서비스 상태(uptime) 페이지.")
b("최종 목표: AWS 이전 — RDS, 컨테이너(ECS 등), S3+CloudFront, Terraform, GitHub Actions CI/CD.")

h1("6. 재개 방법")
p("다음 세션에서 아래 순서로 서버를 띄우면 바로 이어서 작업할 수 있다.")
b('컨테이너: docker compose up -d db mailpit')
b("백엔드: backend/.venv/bin/uvicorn app.main:app --app-dir backend --port 8000")
b("프론트: cd frontend && npm run dev (http://localhost:5173)")
p("데모 계정: kim@test.com / secret123 (글 소유), lee@test.com / pw12345 (구독자)")

doc.save(OUT)
print("저장됨:", OUT)
