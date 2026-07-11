"""2026-07-11 개발일지(일주일 만에 복귀 → 코드 하이라이팅 + 태그로 블로그 채우기 + SQLi 실측 방어)를 Word로.
   세션 복구, 코드블록 색칠, 태그/필터, SQL 인젝션 직접 찔러보기 + GIN 인덱스. WSL + Windows host 양쪽 저장."""

import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DATE = "2026-07-11"
FILENAME = f"블로그_개발일지_{DATE}.docx"
PROJECT_PATH = f"/home/es0764/blog-platform/{FILENAME}"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7E, 0x34)


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def p(t, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.bold = bold
    return para


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def analogy(t):
    para = doc.add_paragraph()
    r = para.add_run("🔎 비유   ")
    r.bold = True
    r.font.color.rgb = BLUE
    r2 = para.add_run(t)
    r2.font.color.rgb = BLUE
    para.paragraph_format.left_indent = Pt(12)


def field(label, text, color=GRAY):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    r = para.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = color
    para.add_run(text)
    return para


def note(t):
    para = doc.add_paragraph()
    r = para.add_run("🛠 전문가 노트   ")
    r.bold = True
    r.font.color.rgb = GREEN
    r2 = para.add_run(t)
    r2.font.color.rgb = GREEN
    para.paragraph_format.left_indent = Pt(12)


# ===================== 표지 =====================
doc.add_heading("블로그 개발일지", level=0)
p(f"날짜: {DATE}", bold=True)
p("주제: 일주일 만에 복귀 — 어디까지 했는지 되짚고, 코드 하이라이팅·태그로 블로그를 더 채우고, SQL 인젝션을 직접 찔러 방어를 확인하기")
p("대상: 웹/클라우드 입문자가 읽어도 이해되게 — 세션 복구, 코드 색칠, 태그 분류, 그리고 'SQL 인젝션'이 뭔지·왜 안 통하는지 위주로.")
doc.add_paragraph("")

# ===================== 0. 한눈에 =====================
h1("0. 오늘 한 일 한눈에")
p("일주일 쉬고 와서 '내가 뭐 하려 했더라'가 기억이 안 났다. 그래서 먼저 기록(PROGRESS)으로 위치를 되짚고, "
  "하던 흐름대로 블로그를 더 채웠다. 코드 블록에 색을 입히고(하이라이팅), 글에 태그를 달아 분류/필터가 되게 했다. "
  "마지막으로 태그 검색 기능에 'SQL 인젝션'을 직접 넣어보며 안전한지 확인하고, 검색 속도를 위한 인덱스도 넣었다.")
bullet("복귀: 기록으로 '어디까지 했는지' 되짚어 위치 복구")
bullet("코드 하이라이팅: 글의 코드 블록에 문법 색칠 (기술 블로그 필수)")
bullet("태그/카테고리: 글에 태그 → 카드/사이드바 표시 + 태그로 필터")
bullet("보안: 태그 검색에 SQL 인젝션 실제로 넣어봄 → 안 통함(파라미터화) 확인")
bullet("성능: 태그 검색이 빠르도록 GIN 인덱스 추가")
analogy("도서관에 '책 표지'(커버)를 이미 달았고, 오늘은 '분류 라벨'(태그)과 '색인'(인덱스)을 붙이고, "
        "'가짜 대출증으로 못 들어오나' 직접 시험해 본 날이다.")

# ===================== 1. 복귀 =====================
h1("1. 복귀 — '내가 뭐 하려 했더라'")
p("일주일 쉬면 맥락이 날아간다. 그래서 이 프로젝트는 모든 굵직한 작업을 PROGRESS 기록과 개발일지로 남긴다. "
  "오늘도 그 기록을 펴서 '여기까지 했고 다음은 이거'를 3분 만에 복구했다.")
field("복구한 위치", "블로그는 이미 라이브(AWS)로 돌고 있고, 최근에 '초라한 블로그를 풍성하게' + '보안(DoS 등)'을 "
                "끝냈으며, 다음 후보는 '태그/코드 하이라이팅/콘텐츠 채우기'였다.", GRAY)
analogy("여행 중 일지를 써두면, 며칠 쉬어도 '어제 어디까지 걸었지?'를 일지 한 줄로 안다. 코드도 똑같다 — "
        "기록이 미래의 나에게 남기는 지도다.")
note("혼자 하는 프로젝트일수록 '기록'이 실력이다. 며칠만 지나도 결정의 이유·다음 할 일이 흐려진다. "
     "PROGRESS/일지는 남 보여주기용이 아니라, 미래의 내가 빨리 이어서 하기 위한 장치다.")

# ===================== 2. 코드 하이라이팅 =====================
h1("2. 코드 하이라이팅 — 코드에 색 입히기")
p("기술 블로그인데 코드 블록이 흑백 고정폭 글씨로만 나왔다. 여기에 문법 색칠(키워드·문자열·주석 색)을 붙였다.")
field("어떻게", "글을 화면에 그리는 도구(react-markdown)에 하이라이팅 플러그인(rehype-highlight)을 끼웠다. "
             "코드의 각 조각을 인식해 색 표시를 붙여준다. 색 테마(라이트/다크)는 CSS로.", GREEN)
field("검증", "python(`def`가 키워드 색), bash, json을 넣어 실제로 색 태그가 생기는지 확인했다.", GREEN)
analogy("워드에서 코드에 색을 칠하듯, 브라우저가 코드를 '단어 종류별로' 알아보고 색을 입히는 것. "
        "읽는 사람이 코드 구조를 한눈에 파악한다.")
note("이건 '독자의 브라우저'에서 색을 칠한다(서버 부담 없음). 대신 색칠 도구(highlight.js)가 통째로 앱에 실려 "
     "용량이 조금 커졌다. 필요하면 '자주 쓰는 언어만' 싣거나, 코드 화면에서만 불러오도록(코드분할) 줄일 수 있다.")

# ===================== 3. 태그 =====================
h1("3. 태그/카테고리 — 분류로 채우기")
p("글에 태그(예: AWS, Terraform)를 달고, 그 태그로 글을 모아 볼 수 있게 했다. 네이버 카페 '게시판' 같은 느낌.")
field("글쓰기", "태그 입력칸에서 치고 Enter를 누르면 '칩'으로 추가되고, ×로 뺀다. 최대 10개.", GREEN)
field("보기/필터", "글 카드·상세에 태그 칩이 뜨고, 사이드바엔 '태그 목록(개수)'이 있다. 태그를 누르면 주소에 "
               "'?tag=AWS'가 붙으며 그 태그 글만 모여 보이고, '전체보기'로 해제한다.", GREEN)
field("저장 방식", "태그는 DB에 '배열'로 저장한다. 글 하나에 여러 태그가 리스트로 들어간다. 서버가 공백정리·중복제거·"
               "개수/길이 제한을 해서 지저분한 입력을 막는다.", GREEN)
analogy("책에 여러 개의 분류 스티커를 붙이고, 스티커 하나를 누르면 같은 스티커가 붙은 책만 서가에 모이는 것.")
note("중요한 함정 하나: 태그 필터는 '공개범위 조건과 반드시 함께(AND)' 걸어야 한다. 안 그러면 비로그인 사용자가 "
     "'?tag=비밀'로 남의 비공개 글을 엿볼 수 있다(권한 우회). 그래서 '이 태그 + 볼 수 있는 글'만 나오게 했다.")

# ===================== 4. 보안 =====================
h1("4. 태그 검색을 직접 털어보기 — SQL 인젝션 + 인덱스")
p("검색어(태그)는 '사용자 입력'이라, 이걸 DB 질의에 잘못 끼우면 'SQL 인젝션'이라는 고전적 공격이 통할 수 있다. "
  "그래서 실제로 공격 문자열을 넣어봤다.")
field("SQL 인젝션이란", "검색어에 DB 명령을 섞어 넣어, 원래 질의를 조작하는 것. 예: 태그에 \"' OR '1'='1\"을 넣어 "
                   "'모든 글'을 빼내거나, \"'); DROP TABLE posts;--\"로 '표 자체를 삭제'하려는 시도.", RED)
field("실측 결과", "세 가지 공격 문자열을 태그 검색에 넣었더니 → 전부 '그런 태그 없음(0개)'으로 처리되고, 글 표도 "
               "멀쩡했다. 우리 도구(SQLAlchemy)가 검색어를 'DB 명령'이 아니라 '그냥 값'으로만 취급(파라미터화)하기 "
               "때문이다. 인젝션 불가.", GREEN)
analogy("택배 송장에 '주소' 대신 '이 집 금고 열어'라고 적어도, 기사(DB)는 그걸 '주소 글자'로만 읽지 명령으로 "
        "실행하지 않는다. 파라미터화가 바로 이 '글자로만 취급' 장치다.")
field("검색 속도(인덱스)", "태그가 배열이라, 그냥 두면 '모든 글을 하나씩 뒤져' 태그를 찾는다(글이 많아지면 느림). "
                    "그래서 배열 검색용 'GIN 인덱스'를 붙여 색인으로 바로 찾게 했다.", GREEN)
note("사용자 입력이 DB로 갈 땐 '문자열을 직접 이어붙이지 말고 파라미터로 넘긴다'가 철칙이다. 요즘 ORM(SQLAlchemy)은 "
     "기본이 파라미터화라 안전하지만, 직접 확인(공격 문자열 넣어보기)해 두면 확실하다. 그리고 배열/전문 검색은 "
     "일반 인덱스가 안 먹으니 GIN 같은 특수 인덱스를 쓴다.")

# ===================== 5. 개념 =====================
h1("5. 오늘 익힌 핵심 개념")
bullet("기록이 실력: 며칠 쉬어도 일지 한 줄로 위치 복구 — 미래의 나에게 남기는 지도")
bullet("하이라이팅은 클라이언트에서: 서버 부담 없이 독자 브라우저가 코드에 색을 칠한다(대신 용량↑)")
bullet("태그 필터는 공개범위와 AND: 안 그러면 태그로 남의 비공개 글이 새는 권한 우회")
bullet("SQL 인젝션 방어 = 파라미터화: 사용자 입력을 'DB 명령'이 아니라 '값'으로만 취급")
bullet("직접 찔러보기: '아마 안전하겠지'가 아니라 공격 문자열을 실제로 넣어 0개·표 멀쩡을 확인")
bullet("특수 인덱스(GIN): 배열/전문 검색은 일반 인덱스가 아니라 GIN으로 색인 스캔")

# ===================== 6. 다음 =====================
h1("6. 다음 할 일")
bullet("(선택) 진짜 글 콘텐츠 채우기 — 코드블록+태그 넣으면 하이라이팅·태그 UI가 다 살아남 (실제 작업이 좋은 글감)")
bullet("(선택) 데모로 넣은 커버 3개를 진짜 이미지로 교체")
bullet("(선택) 남은 낮은 항목: 상태점검 메일연결을 백그라운드 캐시로 / 글 목록은 발췌만 반환")
bullet("(선택) 커스텀 도메인 + 오리진 HTTPS")

# ===================== 저장 =====================
doc.save(PROJECT_PATH)
print(f"[저장] 프로젝트(WSL): {PROJECT_PATH}")

host_user = "/mnt/c/Users/erert"
cands = [
    os.path.join(host_user, "Desktop"),
    os.path.join(host_user, "OneDrive", "Desktop"),
    os.path.join(host_user, "바탕 화면"),
    os.path.join(host_user, "Documents"),
    host_user,
]
host_dir = next((d for d in cands if os.path.isdir(d)), host_user)
host_path = os.path.join(host_dir, FILENAME)
shutil.copy2(PROJECT_PATH, host_path)
print(f"[저장] Windows host: {host_path}")
