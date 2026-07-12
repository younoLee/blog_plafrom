"""2026-07-12 개발일지(남은 DoS 마무리 + 관리자 인프라 대시보드 + WAF/비용 학습)를 Word로.
   status 캐시, 목록 발췌, 관리자 실시간 대시보드, 그리고 'WAF는 왜 못 떼나' 비용 교훈. WSL + Windows host 양쪽 저장."""

import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DATE = "2026-07-12"
FILENAME = f"블로그_개발일지_{DATE}.docx"
PROJECT_PATH = f"/home/es0764/blog-platform/{FILENAME}"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7E, 0x34)


def h1(t):
    doc.add_heading(t, level=1)


def p(t, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.bold = bold
    return para


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def analogy(t):
    para = doc.add_paragraph()
    r = para.add_run("🔎 비유   ")
    r.bold = True
    r.font.color.rgb = BLUE
    r2 = para.add_run(t)
    r2.font.color.rgb = BLUE
    para.paragraph_format.left_indent = Pt(12)


def field(label, text, color=GRAY):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    r = para.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = color
    para.add_run(text)
    return para


def note(t):
    para = doc.add_paragraph()
    r = para.add_run("🛠 전문가 노트   ")
    r.bold = True
    r.font.color.rgb = GREEN
    r2 = para.add_run(t)
    r2.font.color.rgb = GREEN
    para.paragraph_format.left_indent = Pt(12)


# ===================== 표지 =====================
doc.add_heading("블로그 개발일지", level=0)
p(f"날짜: {DATE}", bold=True)
p("주제: 남은 DoS 구멍 마무리 + 관리자 전용 '인프라 상태 대시보드' 만들기, 그리고 'WAF를 왜 못 떼는가'로 배운 클라우드 비용")
p("대상: 웹/클라우드 입문자가 읽어도 이해되게 — 캐시로 부하 줄이기, 응답 가볍게, 서버 혼잡도 보기, 클라우드 요금제 함정 위주로.")
doc.add_paragraph("")

# ===================== 0. 한눈에 =====================
h1("0. 오늘 한 일 한눈에")
p("남아 있던 자잘한 서버 부하 두 개를 정리하고(상태점검 캐시, 목록 응답 가볍게), 관리자만 보는 '서버 혼잡도 대시보드'를 "
  "만들었다. 마지막으로 비용을 줄이려 WAF를 떼려다 막혔는데, 그 과정에서 '클라우드 요금제'의 중요한 걸 배웠다.")
bullet("상태점검 캐시: /status가 호출마다 메일서버에 2초씩 연결하던 걸 '1분마다 갱신되는 캐시'로")
bullet("목록 가볍게: 글 목록이 매번 모든 글 본문을 통째로 보내던 걸 '발췌+읽기시간'만 보내게")
bullet("관리자 대시보드: 서버 CPU·메모리·디스크·부하 + DB 커넥션을 색 막대로, 10초마다 갱신(관리자만)")
bullet("WAF/비용 학습: WAF를 떼려다 실패 → 'CloudFront 무료 요금제에 포함된 무료 WAF'라 못 뗀다는 걸 배움")
analogy("가게로 치면 — 손님 올 때마다 창고까지 전화하던 걸 '메모판 보기'로 바꾸고(캐시), 메뉴판을 요약본으로 "
        "갈고, 사장님용 '실시간 혼잡도 모니터'를 달고, 마지막에 '보험 해지하려다 이미 무료 패키지에 껴 있는 걸 알게 된' 날.")

# ===================== 1. 캐시 =====================
h1("1. 상태점검을 '캐시'로 — 매번 안 물어보게")
p("서비스 상태 페이지가 쓰는 /status 주소는 호출될 때마다 '메일 서버 살아있나?' 하고 소켓 연결을 시도했다(최대 2초). "
  "누가 이걸 마구 부르면 서버 일꾼이 2초씩 묶인다.")
field("고친 방법", "어차피 백그라운드가 1분마다 상태를 점검해 기록하고 있었다. 그 결과를 '캐시(메모리에 마지막 값)'로 "
             "저장해두고, /status는 그 캐시만 읽게 했다. 이제 호출마다 메일서버에 연결하지 않는다.", GREEN)
analogy("직원이 손님마다 창고에 전화(2초)해 재고를 묻던 걸, '1분마다 갱신되는 재고 메모판'을 보게 바꾼 것. "
        "메모는 최대 1분 오래됐을 뿐, 훨씬 빠르고 창고 전화가 사라진다.")
note("'매 요청마다 하는 무거운 작업'은 캐시의 1순위 후보다. 값이 조금 오래돼도 되는 것(상태·통계)이면 특히. "
     "여기선 이미 1분마다 도는 백그라운드가 있어서, 그 결과를 재활용하기만 하면 됐다.")

# ===================== 2. 목록 가볍게 =====================
h1("2. 글 목록을 '발췌'만 — 응답 가볍게")
p("글 목록 주소가 매번 '모든 글의 본문 전체'를 담아 보내고 있었다. 글이 길고 많아지면 목록 한 번에 수십만 자가 실린다.")
field("고친 방법", "목록엔 '본문 전체' 대신 '짧은 발췌 + 읽기시간'만 담게 했다(상세 화면에서만 본문 전체). 발췌는 서버가 "
             "마크다운 기호를 벗겨 깔끔하게 만들어 보낸다.", GREEN)
analogy("메뉴판에 요리 '전체 레시피'를 다 적어두던 걸, '한 줄 소개'만 남긴 것. 자세한 건 주문(상세)했을 때만.")
note("이건 성능이자 보안(증폭 방지)이다. 무인증으로 부를 수 있는 목록이 매번 거대한 응답을 만들면, 공격자가 그걸 반복해 "
     "대역폭·CPU를 소모시킬 수 있다. 목록은 '요약', 상세는 '전체'가 기본이다.")

# ===================== 3. 대시보드 =====================
h1("3. 관리자 인프라 대시보드 — 서버 혼잡도 한눈에")
p("관리자만 보는 화면에 '지금 서버가 얼마나 바쁜지'를 실시간으로 보여주는 대시보드를 만들었다.")
field("무엇을", "서버(EC2)의 CPU·메모리·디스크·부하, 그리고 DB 커넥션 수를 색 막대(초록<60% / 노랑<85% / 빨강≥85%)로. "
             "10초마다 자동 갱신되고, 관리자만 볼 수 있다.", GREEN)
field("어떻게(그리고 왜 공짜)", "AWS의 CloudWatch(서비스별 지표)는 호출마다 돈이 든다. 그래서 대신 서버가 '자기 자신의 "
                       "상태(CPU·메모리 등)'를 직접 읽는 도구(psutil)와 DB 질의를 썼다. 새 AWS 리소스 0, 추가비용 0.", GREEN)
analogy("자동차 계기판을 단 것 — 엔진 온도·연료를 실시간 바늘로. 단 '정비소에 전화해 묻기(CloudWatch, 유료)' 대신 "
        "'차 안의 센서를 직접 읽기(psutil, 무료)'로 만들었다.")
note("모니터링은 '어디서 재느냐'로 비용이 갈린다. 관리형 지표(CloudWatch)는 편하지만 과금되고, 서버 자체 지표(psutil)는 "
     "공짜지만 그 서버 한 대만 본다. 단일 인스턴스 취미 프로젝트엔 후자가 '가성비'다.")

# ===================== 4. WAF/비용 =====================
h1("4. WAF를 떼려다 배운 클라우드 비용")
p("월 비용을 줄이려 방화벽(WAF)을 떼려 했다. 그런데 코드로 떼려니 에러가 났다: '요금제 구독이 있는 배포는 WAF가 필수다.'")
field("무슨 일이었나", "이 블로그의 CDN(CloudFront)은 'Free 요금제(flat-rate)'에 들어 있었고, 그 요금제는 WAF를 '번들로 "
                "필수' 포함한다. 즉 WAF는 따로 돈 나가는 게 아니라 무료 패키지의 일부였다. 떼려면 유료 종량제로 바꿔야 "
                "하는데, 그럼 오히려 CDN에 돈이 나가기 시작한다.", RED)
field("결론", "내 'WAF 월 $8' 추정이 틀렸다. WAF는 사실상 무료라 그냥 두는 게 이득. 변경을 원복했다(실패한 시도라 "
            "라이브는 안 건드려졌다).", GREEN)
analogy("보험료 아끼려 '화재보험'을 해지하려 했더니, 알고 보니 이미 가입한 '무료 통신 패키지'에 껴 있어서 따로 돈이 안 "
        "나가던 것. 해지하면 오히려 패키지 혜택이 깨진다.")
note("클라우드 비용은 '겉보기'와 다를 때가 많다. 리소스 하나의 가격만 보지 말고 '어떤 요금제/번들에 묶여 있는가'를 봐야 "
     "한다. 그리고 추정은 늘 실제 청구/문서로 검증할 것 — 나도 이번에 추정이 틀렸다.")

# ===================== 5. 비용 정리 =====================
h1("5. 클라우드 비용, 지금 정리")
p("프리티어(12개월)가 끝나도 24시간 켜두면 대략 월 $33 수준이고, 대부분이 DB와 서버다.")
bullet("RDS(관리형 DB) ~$21/월 — 제일 큼")
bullet("EC2(서버) ~$12/월")
bullet("CloudFront + WAF ~$0 — Free 요금제 번들(1TB/월 무료, 상시)")
bullet("가장 큰 절약: 안 쓸 때 EC2·RDS 정지(→거의 $0) / DB를 서버 안으로 옮기기(-$21)")
p("그래서 오늘은 끝내면서 EC2·RDS를 '정지'해 비용을 멈춘다. (다시 켤 때 서버 IP가 바뀌어 CDN 오리진을 갱신해야 함)")
note("취미/포트폴리오 인프라의 정석: '안 쓸 땐 끈다.' 정지 중엔 컴퓨팅 요금이 0에 가깝다. 대신 켤 때마다 바뀌는 IP를 "
     "감안해 자동화(오리진 갱신)를 해두면 편하다.")

# ===================== 6. 다음 =====================
h1("6. 다음 할 일")
bullet("(다음 세션) EC2·RDS 다시 켜기 + CloudFront 오리진(새 EC2 주소) 갱신")
bullet("(선택) 진짜 글 콘텐츠 채우기 — 코드블록+태그로 하이라이팅·태그가 다 살아남")
bullet("(선택) 비용 더: RDS를 EC2 안 Postgres로 이전 / 예산 알림(AWS Budgets)")
bullet("(선택) CloudWatch로 AWS 서비스별 상세 지표 확장(유료), 커스텀 도메인")

# ===================== 저장 =====================
doc.save(PROJECT_PATH)
print(f"[저장] 프로젝트(WSL): {PROJECT_PATH}")

host_user = "/mnt/c/Users/erert"
cands = [
    os.path.join(host_user, "Desktop"),
    os.path.join(host_user, "OneDrive", "Desktop"),
    os.path.join(host_user, "바탕 화면"),
    os.path.join(host_user, "Documents"),
    host_user,
]
host_dir = next((d for d in cands if os.path.isdir(d)), host_user)
host_path = os.path.join(host_dir, FILENAME)
shutil.copy2(PROJECT_PATH, host_path)
print(f"[저장] Windows host: {host_path}")
