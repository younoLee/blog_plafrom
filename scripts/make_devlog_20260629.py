"""2026-06-29 개발일지(디스코드 검은화면 버그 진단·해결 + terraform 드리프트 + 배포)를 Word로.
   '추측 말고 측정'하는 디버깅 과정을 입문자도 따라오게. WSL + Windows host 양쪽 저장."""

import os
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DATE = "2026-06-29"
FILENAME = f"블로그_개발일지_{DATE}.docx"
PROJECT_PATH = f"/home/es0764/blog-platform/{FILENAME}"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
RED = RGBColor(0xC0, 0x39, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x7E, 0x34)


def h1(t):
    doc.add_heading(t, level=1)


def h2(t):
    doc.add_heading(t, level=2)


def p(t, bold=False):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.bold = bold
    return para


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def analogy(t):
    para = doc.add_paragraph()
    r = para.add_run("🔎 비유   ")
    r.bold = True
    r.font.color.rgb = BLUE
    r2 = para.add_run(t)
    r2.font.color.rgb = BLUE
    para.paragraph_format.left_indent = Pt(12)


def field(label, text, color=GRAY):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(12)
    r = para.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = color
    para.add_run(text)
    return para


def note(t):
    para = doc.add_paragraph()
    r = para.add_run("🛠 전문가 노트   ")
    r.bold = True
    r.font.color.rgb = GREEN
    r2 = para.add_run(t)
    r2.font.color.rgb = GREEN
    para.paragraph_format.left_indent = Pt(12)


# ===================== 표지 =====================
doc.add_heading("블로그 개발일지", level=0)
p(f"날짜: {DATE}", bold=True)
p("주제: '디스코드에서 AI 초안 누르면 검은 화면' 버그 — 추측 말고 측정으로 원인 찾기 + 해결 + 배포 + terraform 드리프트 정리")
p("대상: 웹/클라우드 입문자가 읽어도 이해되도록 — 디버깅을 '어떻게 생각하며' 푸는지 위주로.")
doc.add_paragraph("")

# ===================== 0. 한눈에 =====================
h1("0. 오늘 한 일 한눈에")
p("어제 만든 AI 초안 기능에서 버그 신고가 왔다 — '모바일/디스코드에서 초안을 누르면 검은 화면이 뜨고 안 된다.' "
  "오늘은 이걸 추측으로 코드를 갈아엎지 않고, 가설 → 측정 → 확인 순서로 원인을 좁혀 해결했다. "
  "겸사겸사 terraform(인프라 코드)에서 튀어나온 별개 에러도 정리했다.")
bullet("증상 신고: 디스코드에서 AI 초안 → 전체 검은 화면, 새로고침해야 복구")
bullet("진단: '크래시'가 아니라 '너무 오래 걸려서'였다 — 생성에 47초")
bullet("해결: 기본 모델을 빠른 Haiku로 (47초 → 10초) + 안전장치들")
bullet("배포 함정: '다 했어'가 알고 보니 절반만 — 검증으로 잡아냄")
bullet("덤: terraform이 RDS 버전을 되돌리려다 낸 에러 정리")
analogy("어제가 'AI 비서를 채용한 날'이라면, 오늘은 '그 비서가 특정 방(디스코드)에서만 먹통이 되는 이유를 추적해, "
        "일을 더 빨리 하도록 바꿔준 날'이다.")

# ===================== 1. 진단 =====================
h1("1. 버그 진단 — '추측 말고 측정'")
p("가장 중요한 교훈이 여기 있다. '검은 화면'이라고 바로 코드를 의심하지 않고, 단서를 하나씩 모아 범위를 좁혔다.")

h2("단서 1 — '검은 화면 + 새로고침해야 복구'")
p("이건 보통 화면을 그리는 코드(React)가 예외로 죽어서 화면 전체가 사라진 것이다. 그래서 처음엔 '렌더 크래시'를 의심했다.")
analogy("무대(화면) 위 배우가 갑자기 쓰러지면 무대가 통째로 깜깜해지는 것 — 처음엔 '대본(코드)에 문제가 있나' 의심.")

h2("단서 2 — '크롬에서는 기다리면 된다'")
field("핵심 반전", "사용자에게 '폰 크롬에서도 그런가?' 물으니 '그건 기다리면 뜬다'고 했다. 크롬에서 멀쩡히 뜬다는 건 "
                "코드(대본)는 멀쩡하다는 뜻이다. 즉 크래시가 아니다. 문제는 '오래 걸린다'는 것이고, "
                "디스코드 인앱 브라우저만 그 긴 시간을 못 버티는 것이다.", GREEN)
note("디스코드/인스타 안에서 링크를 누르면 뜨는 화면은 진짜 브라우저가 아니라 '웹뷰(WebView)'라는 축소판이다. "
     "웹뷰는 30초 넘게 걸리는 요청에서 잘 멈춘다. 그래서 같은 페이지라도 크롬은 OK, 웹뷰는 검은 화면.")

h2("단서 3 — 시간을 직접 쟀다 (측정)")
field("측정", "추측을 멈추고 운영 서버에서 실제 생성 시간을 쟀다. 같은 메모로 기본 모델(Claude Sonnet)이 "
            "47.2초가 걸렸다. 의심이 사실로 바뀌는 순간이다.", GREEN)
analogy("'느린 것 같다'는 느낌을 스톱워치로 '47초'라는 숫자로 만든 것. 숫자가 있으면 더 못 우긴다.")

h2("곁가지 — 틀린 가설도 측정으로 버린다")
p("처음엔 'CloudFront(우리 사이트의 정문 역할)가 30초만 기다리고 끊어서 그런가?'라고 의심했다. 그런데 설정을 직접 확인하니 "
  "이미 60초로 되어 있었다. 47초는 60초 안이라 끊기지 않는다. 그래서 이 가설은 버렸다 — 크롬이 47초를 기다려준 이유와도 들어맞는다.")
note("틀린 가설을 '측정'으로 빨리 버리는 것도 디버깅의 핵심이다. 추측을 붙들고 코드를 고치면 엉뚱한 데를 파게 된다.")

# ===================== 2. 해결 =====================
h1("2. 해결 — 빠르게 만들면 웹뷰도 버틴다")
p("원인이 '시간'이니, 해결도 '시간 줄이기'다. 가장 큰 레버는 모델 교체였다.")
field("① 기본 모델 Sonnet → Haiku", "Claude에는 빠른 모델(Haiku)·균형(Sonnet)·고품질(Opus)이 있다. 초안 '뼈대 잡기'엔 "
                              "Haiku로 충분하고 2~3배 빠르다. 같은 메모가 47초 → 16초(나중 실측 10초)로 줄었다. "
                              "기본값만 바꾼 거라, 품질이 필요하면 글쓰기 화면 드롭다운에서 Sonnet/Opus를 그대로 고를 수 있다.", GREEN)
field("② 생성 분량 상한 낮춤", "한 번에 만들 최대 분량(MAX_TOKENS)을 4000 → 2500으로. 긴 초안이 끝없이 길어지는 걸 막아 "
                        "시간·비용을 동시에 줄인다.", GREEN)
field("③ 클라이언트 안전장치", "화면 쪽에 90초 타임아웃을 걸어, 응답이 영영 안 와도 무한 대기 대신 '오래 걸려서 멈췄어' 메시지로 끝낸다.", GREEN)
field("④ 에러 경계(ErrorBoundary)", "혹시 진짜로 화면 코드가 죽어도 '검은 화면' 대신 에러 메시지를 띄우고 앱 전체가 안 죽게 했다. "
                            "(이번 원인은 아니었지만, 앱이라면 갖춰야 할 안전망)", GREEN)
field("⑤ 사용자 안내", "글쓰기 화면에 '생성에 길면 1분, 디스코드 같은 앱 안 브라우저에선 멈출 수 있으니 크롬에서 써줘' 안내를 추가.", GREEN)
analogy("배달이 너무 오래 걸려 손님(웹뷰)이 가버리는 가게 → '더 빠른 요리사(Haiku)'로 바꿔 10분 안에 내보내는 것. "
        "정 고급 요리가 필요하면 손님이 '느려도 괜찮은 자리(크롬)'에서 주문하면 된다.")

# ===================== 3. 배포 함정 =====================
h1("3. 배포 함정 — '다 했어'가 절반이었다")
p("수정을 운영에 올려야 효과가 난다. 프론트(화면)는 git push로 자동 배포, 백엔드(서버 로직)는 코드를 압축해 서버로 보내 재빌드한다.")
field("무슨 일", "사용자가 '다 했어' 해서 검증해보니, 프론트만 올라갔고 백엔드는 옛날 코드 그대로(컨테이너가 23시간째 안 바뀜, "
            "여전히 Sonnet·4000)였다. 즉 진짜 해결책(Haiku)이 라이브에 없었다.", RED)
field("해결", "내가 백엔드 코드를 압축·전송하고(이건 단순 작업이라 내가), 사용자가 서버에서 재빌드 명령을 실행했다. "
            "그 뒤 컨테이너 안의 값이 Haiku·2500으로 바뀐 걸 확인했다.", GREEN)
field("검증(끝까지)", "운영 사이트에 실제 요청을 보내 기본 초안 생성이 10.2초·정상(200)·코드 없음인 걸 확인. "
                "47초 → 10초로 실측 단축. 이제 디스코드 웹뷰도 버틸 수준.", GREEN)
analogy("'설치 다 했어요'라고 했는데 실제로 켜보니 절반만 된 것. 그래서 '말'이 아니라 '실제 동작'으로 확인해야 한다.")
note("배포 후엔 항상 '돌아가는 실물'을 확인하는 습관이 안전하다. 코드를 서버에 복사했다고 끝이 아니라, "
     "재빌드해서 '실행 중인 것'이 바뀌어야 진짜 반영이다. (디스크의 새 코드 ≠ 실행 중 컨테이너)")

# ===================== 4. terraform 드리프트 =====================
h1("4. 덤 — terraform 'RDS 버전' 에러")
p("인프라 코드(terraform)로 apply하려다 에러가 났다. AI 버그와는 무관한, 따로 쌓여 있던 문제였다.")
field("원인", "AWS가 데이터베이스(RDS)의 소소한 버전을 자동으로 16.12 → 16.13으로 올려놨는데, 우리 코드엔 아직 16.12로 적혀 있었다. "
            "terraform이 '코드대로 16.12로 되돌리자'고 했지만, DB 버전은 내릴 수가 없어서 AWS가 거부 → 에러.", RED)
field("해결", "코드를 현실(16.13)에 맞추고, 'engine_version은 terraform이 건드리지 말고 AWS 자동 업그레이드에 맡겨라'(ignore_changes)고 "
            "지정했다. 다시 확인하니 '바뀔 것 없음(No changes)'으로 깨끗해졌다.", GREEN)
analogy("재고 장부(코드)엔 옛 모델이라 적혀 있는데 창고(실제 AWS)엔 신형이 들어와 있는 상황. "
        "장부를 현실에 맞추고, '이 항목은 자동 갱신되니 장부가 따지지 마라'고 표시한 것.")
note("드리프트(drift) = 코드와 실제 인프라가 어긋난 상태. 자동 업그레이드처럼 '내가 안 했는데 바뀌는' 항목은 "
     "ignore_changes로 빼두면 매번 싸우지 않는다.")

# ===================== 5. 개념 =====================
h1("5. 오늘 익힌 핵심 개념")
bullet("디버깅은 추측이 아니라 측정 — '느린 것 같다'를 '47초'로 만든다. 틀린 가설(30초 컷)도 측정으로 빨리 버린다")
bullet("증상의 범위를 좁혀라 — '크롬은 됨 / 디스코드는 안 됨' 한 줄이 '코드 버그 아님 + 웹뷰 한계'로 원인을 확정")
bullet("인앱 브라우저(웹뷰)는 진짜 브라우저가 아니다 — 긴 요청에 약하다. 무거운 작업은 일반 브라우저 권장")
bullet("성능은 모델 선택의 문제 — 작업에 맞는 가장 빠른 모델(Haiku)을 기본으로, 품질이 필요할 때만 큰 모델")
bullet("배포는 '말'이 아니라 '실물'로 검증 — 디스크의 새 코드 ≠ 실행 중 컨테이너. 끝까지 확인")
bullet("인프라 드리프트 — 자동으로 바뀌는 항목은 ignore_changes로 빼서 코드와 현실의 싸움을 막는다")

# ===================== 6. 다음 =====================
h1("6. 다음 할 일")
bullet("(선택) 진짜 끝장 해결 = 응답 스트리밍 — 생성되는 대로 조금씩 보내면 웹뷰도 '살아있음' 인식. 백엔드/CloudFront/프론트 다 손대는 큰 작업")
bullet("(선택) 모델별 예상 소요시간을 드롭다운에 표시(예: Haiku 빠름 / Sonnet 느림)")
bullet("(선택) 다음 로드맵: 도메인 + SES 메일, 백엔드 자동배포(SSH), 상태 페이지 고도화")

# ===================== 저장 =====================
doc.save(PROJECT_PATH)
print(f"[저장] 프로젝트(WSL): {PROJECT_PATH}")

host_user = "/mnt/c/Users/erert"
cands = [
    os.path.join(host_user, "Desktop"),
    os.path.join(host_user, "OneDrive", "Desktop"),
    os.path.join(host_user, "바탕 화면"),
    os.path.join(host_user, "Documents"),
    host_user,
]
host_dir = next((d for d in cands if os.path.isdir(d)), host_user)
host_path = os.path.join(host_dir, FILENAME)
shutil.copy2(PROJECT_PATH, host_path)
print(f"[저장] Windows host: {host_path}")
