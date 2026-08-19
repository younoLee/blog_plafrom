"""개발일지(.docx)를 블로그 글 마크다운으로 변환.

개발일지는 make_devlog_*.py가 python-docx로 만든 것이라 문단 구조가 일정하다.
그 구조를 마크다운으로 되돌린다:

  Title "블로그 개발일지"        → 버림(제목은 POSTS 맵에서 지정)
  표지의 작성일/스택/주제 줄     → 버림(제목·태그로 대체). '대상:'만 인용문 리드로 남김
  Heading 1 "3. 관리자 대시보드" → "## 3. 관리자 대시보드"  (Heading 2 → ###)
  List Bullet                    → "- "
  🔎 비유 / 🛠 전문가 노트        → 인용문(>)
  field("라벨", "내용")          → "**라벨** — 내용"
  그 외 문단                     → 그대로

제목·태그를 자동 추출하지 않고 POSTS에 손으로 적는 이유: 초기 3편(06-21·06-22·06-24)은
표지 형식이 달라 '주제:' 줄이 없고, 있는 편들도 주제 줄이 제목으로 쓰기엔 너무 길다.

사용:
  python scripts/devlog_to_markdown.py                 # 전체 → content/devlog/*.md
  python scripts/devlog_to_markdown.py 2026-07-12      # 한 편만 stdout으로 (미리보기)
  python scripts/devlog_to_markdown.py 2026-07-12 --write   # 한 편만 파일로 (새 편 발행용)

⚠️ 인자 없이 돌리면 **발행된 편을 전부 다시 쓴다.** 2026-08-18에 33편의 em 대시를
손으로 정리했으므로 그게 통째로 되돌아간다. 새 편을 낼 때는 `--write`를 쓸 것.
"""

import re
import sys
from pathlib import Path

from docx import Document


# 제목·태그 표는 devlog_posts.py로 뺐다 — docx 없이 읽을 수 있어야 하고(로컬엔 pip이
# 없다) 프론트 빌드도 태그가 필요하기 때문이다. 이름은 그대로 쓴다.
from devlog_posts import OUT_DIR, POSTS, ROOT, write_tags_json  # noqa: E402


# 본문에서 이 접두사로 시작하는 문단은 인용문으로 뽑는다.
CALLOUTS = ("🔎 비유", "🛠 전문가 노트")

# 표지에서 버릴 메타 줄(제목·태그가 대신한다).
DROP_META = ("작성일:", "날짜:", "스택:", "목표:", "주제:", "오늘 주제:", "오늘의 주제:")


def _is_mono(para) -> bool:
    """터미널 출력 문단인가 — 실행이 하나라도 고정폭 글꼴이면 그렇게 본다.

    make_devlog의 ev()가 Consolas로 찍는다. 캡션("▶ 실제 출력 — …")은 굵은 본문 글꼴이라
    여기 안 걸리고, 코드블록 위에 라벨로 남는다.
    """
    return any(r.font.name == "Consolas" for r in para.runs)


def _is_bullet(para) -> bool:
    return para.style.name in ("List Bullet", "List Bullet 2", "List Number")


def _heading_level(para) -> int | None:
    """Heading 1 → 1. 제목(Title)은 0. 본문이면 None."""
    if para.style.name in ("Title", "Heading 0"):
        return 0
    m = re.match(r"^Heading (\d+)$", para.style.name)
    return int(m.group(1)) if m else None


def _callout(text: str) -> str | None:
    for mark in CALLOUTS:
        if text.startswith(mark):
            return f"> **{mark}** {text[len(mark):].strip()}"
    return None


def _field(para) -> str | None:
    """field()로 만든 문단(첫 run이 bold 라벨 + 나머지가 내용)이면 '**라벨** — 내용'.

    문단 전체가 bold면 단순 강조 문단이므로 field가 아니다.
    """
    runs = [r for r in para.runs if r.text.strip()]
    if len(runs) < 2 or not runs[0].bold or all(r.bold for r in runs):
        return None
    label = runs[0].text.strip()
    body = "".join(r.text for r in runs[1:]).strip()
    return f"**{label}** — {body}" if label and body else None


def _cover_line(text: str) -> str | None:
    """표지 문단 → 남길 마크다운(없으면 None).

    '대상:'은 이 글이 누구를 위한 글인지라 리드 인용문으로 남기고,
    나머지 메타(작성일·스택·주제…)와 '2026-06-22 · 스택: …' 형태는 버린다.
    """
    if text.startswith("대상:"):
        return f"> {text[3:].strip()}"
    if any(text.startswith(k) for k in DROP_META):
        return None
    if re.match(r"^\d{4}-\d{2}-\d{2}\s*[·・]", text):  # "2026-06-22 · 스택: …"
        return None
    return text  # 그 외 표지 문단(리드 설명 등)은 본문으로 살린다


# ── 스타일이 평탄화된 회차 되살리기 ────────────────────────────────────────
# 2026-07-24 편은 Heading·List Bullet이 전부 Normal로 저장돼 있다(같은 make_devlog가
# add_heading을 쓰는데도 그렇다 — 파일이 한 번 다른 도구를 거친 것으로 보인다).
# 그대로 변환하면 8개 절이 전부 사라져 제목 1개짜리 마크다운이 나온다.
#
# 되살리는 조건을 **문서 단위**로 건다: "이 문서에 Heading 스타일이 하나도 없을 때만".
# 문단 단위로 걸면 스타일이 멀쩡한 회차에서 본문 중 "3. 세 번째 이유는 …" 같은 문장이
# 절 제목으로 잘못 승격된다. 평탄화된 문서는 어차피 잃을 구조가 없으니 안전하다.
FLAT_HEADING = re.compile(r"^(\d+)\.(\d+)?\s+\S")
FLAT_BULLET = re.compile(r"^[•·]\s+")
FLAT_TITLE = "블로그 개발일지"  # Title 스타일이 날아가 본문처럼 남은 표지 제목

# 절 제목은 짧고 마침표로 끝나지 않는다. 어미로는 못 거른다 — 이 회차의 절 제목
# 8개 중 5개가 "…했다"로 끝난다("1. 시작이 좋지 않았다 — 아침에 저장소가 깨져 있었다").
FLAT_HEADING_MAX = 80

# 이 워크어라운드를 적용할 회차를 **명시적으로 적는다**. 새 회차가 평탄화된 채로 들어오면
# 조용히 휴리스틱을 태우는 대신 멈춘다.
#
# 왜 자동으로 안 하는가 — 이 휴리스틱은 07-24 문서 하나를 실제로 열어보고 맞춘 것이다.
# "숫자.으로 시작하고 80자 이하면 절 제목", "문단 안 줄바꿈은 전부 정렬이 의미를 갖는
# 덩어리"는 그 문서에서 참인 것이지 일반 규칙이 아니다(그 문서에선 해당 문단 3개가
# 구성도·터미널 출력 전부였다 — 세어서 확인했다). 다른 문서에 그대로 먹이면 인용문이
# 코드블록이 되거나 본문 문장이 절 제목으로 승격되는데, **에러 없이** 그렇게 된다.
# 이 저장소가 반복해서 당한 게 정확히 그 '조용한 실패'다.
FLAT_ALLOWED = frozenset({"2026-07-24"})


def _is_flattened(doc) -> bool:
    return not any(
        (lv := _heading_level(p)) is not None and lv >= 1 for p in doc.paragraphs
    )


def _flat_heading_level(text: str) -> int | None:
    m = FLAT_HEADING.match(text)
    if not m or len(text) > FLAT_HEADING_MAX or text.rstrip().endswith("."):
        return None
    return 2 if m.group(2) else 1


def convert(path: Path) -> tuple[str, str, list[str]]:
    """docx → (제목, 마크다운 본문, 태그)."""
    doc = Document(str(path))
    date = path.stem.split("_")[-1]
    if date not in POSTS:
        raise SystemExit(f"{date}의 제목·태그가 POSTS에 없습니다. 스크립트에 추가하세요.")
    title, tags = POSTS[date]

    blocks: list[str] = []
    # 고정폭(터미널 출력) 줄을 모았다가 한 덩어리로 코드블록에 넣는다. 아래 _flush_mono 참고.
    mono_buf: list[str] = []
    seen_heading = False
    flat = _is_flattened(doc)
    if flat and date not in FLAT_ALLOWED:
        raise SystemExit(
            f"{date}: 이 문서에 Heading 스타일이 하나도 없습니다(평탄화). 그대로 변환하면\n"
            f"  절이 통째로 사라진 껍데기가 나오고, 되살리는 휴리스틱은 2026-07-24 문서\n"
            f"  하나에 맞춰 만든 것이라 다른 문서에는 조용히 잘못 먹을 수 있습니다.\n"
            f"  → 문서를 열어 확인한 뒤, 맞다면 FLAT_ALLOWED에 '{date}'를 추가하세요."
        )

    def _flush_mono() -> None:
        if mono_buf:
            blocks.append("```\n" + "\n".join(mono_buf) + "\n```")
            mono_buf.clear()

    for para in doc.paragraphs:
        text = para.text.rstrip()
        if not text.strip():
            continue

        # 터미널 출력(make_devlog의 ev())은 Consolas 문단으로 들어온다. 예전엔 이걸 못 알아봐서
        # **한 줄이 문단 하나씩** 흩어졌고, 표처럼 열을 맞춘 출력은 정렬이 통째로 사라졌다.
        # 이 회차들의 형식이 "터미널 출력을 그대로 싣는다"인데 발행본에서만 안 지켜지고 있었다.
        # (2026-07-28에 #21을 변환하다 발견. #20도 같은 상태였다)
        if _is_mono(para):
            mono_buf.append(text)
            continue
        _flush_mono()

        if flat and text == FLAT_TITLE:
            continue  # Title 스타일이 날아간 표지 제목

        level = _heading_level(para)
        if level == 0:
            continue  # 표지 제목은 버린다
        if level is None and flat:
            level = _flat_heading_level(text)

        if level is not None:
            seen_heading = True
            blocks.append(f"{'#' * (level + 1)} {text}")
            continue

        if flat and "\n" in text:
            # 문단 안의 줄바꿈(w:br)은 이 회차에선 전부 정렬이 의미를 갖는 덩어리다
            # (표지 구성도 1개 + 터미널 출력 2개). 그냥 두면 마크다운이 줄바꿈을 공백으로
            # 합쳐 한 줄로 뭉갠다. 스타일이 살아 있는 회차는 이 경로를 안 탄다.
            # 표지 영역에도 하나 있어서 _cover_line보다 먼저 본다.
            blocks.append(f"```\n{text}\n```")
            continue

        if not seen_heading:  # 아직 표지 영역
            line = _cover_line(text)
            if line:
                blocks.append(line)
            continue

        callout = _callout(text)
        if callout:
            blocks.append(callout)
            continue

        field = _field(para)
        if field:
            blocks.append(field)
            continue

        flat_bullet = bool(flat and FLAT_BULLET.match(text))
        # 스타일이 날아간 회차는 글머리표가 문단 첫 글자로 남아 있다("•  증명 3종: …").
        bullet_text = FLAT_BULLET.sub("", text) if flat_bullet else text
        if _is_bullet(para) or flat_bullet:
            # 연속된 불릿은 한 블록으로 묶어야 마크다운 목록이 끊기지 않는다
            if blocks and blocks[-1].startswith("- "):
                blocks[-1] += f"\n- {bullet_text}"
            else:
                blocks.append(f"- {bullet_text}")
            continue

        blocks.append(text)

    _flush_mono()  # 문서가 터미널 출력으로 끝나는 경우

    body = "\n\n".join(blocks)
    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    return title, body, tags


def main() -> None:
    files = sorted(ROOT.glob("블로그_개발일지_*.docx"))
    if not files:
        sys.exit("개발일지 .docx를 찾지 못했습니다.")

    args = [a for a in sys.argv[1:] if a != "--write"]
    write = "--write" in sys.argv

    if args:
        target = next((f for f in files if args[0] in f.name), None)
        if not target:
            sys.exit(f"{args[0]} 개발일지를 찾지 못했습니다.")
        title, body, tags = convert(target)
        # `--write`가 없으면 예전처럼 stdout 미리보기다.
        #
        # **왜 한 편만 쓰는 길이 필요했나**: 아래 전체 루프는 발행된 편을 전부 다시 쓰는데,
        # 2026-08-18에 33편의 em 대시를 손으로 정리했으므로 지금 그걸 돌리면 그 작업이
        # 통째로 되돌아간다. 그래서 새 편을 낼 때 쓸 수 있는 명령이 사실상 없었고,
        # 미리보기 출력을 리다이렉트해서 머리말을 손으로 떼는 수밖에 없었다
        # (2026-08-19에 실제로 그렇게 하다가 이 입구를 냈다).
        if not write:
            print(f"제목: {title}\n태그: {tags}\n글자수: {len(body):,}\n{'─' * 60}\n{body}")
            return
        date = target.stem.replace("블로그_개발일지_", "")
        out = OUT_DIR / f"{date}.md"
        out.write_text(f"# {title}\n\n{body}\n", encoding="utf-8")
        print(f"{out}  {len(body):,}자  {title}")
        return

    # ⚠️ **이 루프는 32편을 전부 다시 쓴다 — 손으로 고친 편도 덮는다.**
    #
    # 2026-08-15에 실제로 겪었다. 08-14에 1편·3편의 맨몸 URL과 폐지된 데모계정 주소를
    # 백틱으로 감싸 고쳤는데, 그 수정이 **마크다운에만** 있었다(원본 docx는 그대로다).
    # 새 편을 만들려고 이 스크립트를 돌리자 두 편이 원래대로 되돌아갔다.
    #
    # 지금은 gen-static.mjs의 GFM/autolink 가드가 그걸 잡아 **빌드를 세운다** — 그날도
    # 그렇게 알았다. 즉 조용한 회귀는 아니다. 그래도 알고 돌리는 것과 모르고 돌리는 것은
    # 다르니 여기 적어둔다: **돌린 뒤 `git status content/devlog`를 볼 것.**
    # 오늘 편 말고 다른 편이 바뀌어 있으면 그건 되살아난 옛 원고다(`git checkout`으로 되돌린다).
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for f in files:
        title, body, tags = convert(f)
        date = f.stem.split("_")[-1]
        (OUT_DIR / f"{date}.md").write_text(f"# {title}\n\n{body}\n", encoding="utf-8")
        print(f"{date}  {len(body):>6,}자  {title}")

    # 마크다운을 새로 썼으면 태그 표도 같이 내보낸다. **여기서 같이 하지 않으면**
    # 새 편의 .md는 생겼는데 tags.json은 옛날 것이라, 프론트 빌드가 태그 없는 편을
    # 만나 멈춘다(gen-static.mjs가 그때 실패하도록 해뒀다 — 조용히 빈 태그로
    # 넘어가면 아무도 모른 채 필터에서 빠진다).
    print(f"{write_tags_json()}  ({len(POSTS)}편)")


if __name__ == "__main__":
    main()
