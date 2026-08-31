"""마크다운 문서를 Word(.docx)로 옮긴다 — 설명용 자료처럼 개발일지가 아닌 문서용.

개발일지는 `devlog_md_to_doc.py`가 표지 서식(날짜·시리즈·주제·대상)까지 맞춰 만든다.
이건 그 서식이 필요 없는 일반 문서를 옮기는 쪽이라 따로 뒀다. 글꼴은 같다(맑은 고딕 10.5pt).

지원: #/##/### 제목 · 인용(>) · 글머리표(-) · 번호목록 · 표(|) · 굵게(**) · 코드(`)
목록과 인용은 여러 줄에 걸쳐 있어도 한 항목으로 합친다. 줄마다 문단을 만들면 Word에서
글머리표가 줄 수만큼 생겨 읽기 어렵다.

사용:
    python scripts/md_to_docx.py docs/talk-track-project.md ~/out.docx
    python scripts/md_to_docx.py docs/talk-track-project.md --desktop 블로그_설명자료_프로젝트.docx
"""

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

GRAY = RGBColor(0x55, 0x55, 0x55)
BLUE = RGBColor(0x1F, 0x6F, 0xEB)
MONO = RGBColor(0x24, 0x29, 0x2E)

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(para, text):
    """**굵게** 와 `코드` 를 살려서 문단에 붙인다."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            para.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = para.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = MONO
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            para.add_run(piece)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


BLOCK_START = re.compile(r"^\s*(#{1,3}\s|[-*]\s|>\s|\||\d+\.\s|---\s*$)")


def take_block(lines, i, first):
    """이어지는 줄을 한 덩어리로 먹는다.

    마크다운에서 목록 항목과 인용은 다음 줄로 이어져도 같은 항목이다. 줄 단위로
    문단을 만들면 Word에서 글머리표가 줄마다 하나씩 생겨 읽기 어렵다.
    """
    buf = [first]
    while i + 1 < len(lines) and lines[i + 1].strip() and not BLOCK_START.match(lines[i + 1]):
        i += 1
        buf.append(lines[i].strip())
    return i, " ".join(buf)


def main(src, out, src_label=None):
    src_label = src_label or src
    lines = open(src, encoding="utf-8").read().split("\n")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 표 — 헤더 줄과 구분선을 확인한 뒤 통째로 먹는다
        if stripped.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].strip()) <= set("|-: "):
            header = split_row(stripped)
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                add_runs(cell.paragraphs[0], text)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for c, text in enumerate(row[: len(header)]):
                    cells[c].text = ""
                    add_runs(cells[c].paragraphs[0], text)
            doc.add_paragraph()
            continue

        # 코드/도식 블록 — ``` 사이를 고정폭으로 그대로 옮긴다.
        # 말로 설명하는 자료에서 화살표 도식이 그림 역할을 하므로 서식이 중요하다.
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            for text in block:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(18)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
                run = para.add_run(text if text.strip() else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = MONO
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("> "):
            # 인용은 연속된 > 줄을 한 문단으로 합친다
            buf = [stripped[2:]]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("> "):
                i += 1
                buf.append(lines[i].strip()[2:])
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            add_runs(para, " ".join(buf))
            for r in para.runs:
                r.font.color.rgb = BLUE
                r.italic = True
        elif re.match(r"^\d+\.\s", stripped):
            i, text = take_block(lines, i, re.sub(r"^\d+\.\s", "", stripped))
            add_runs(doc.add_paragraph(style="List Number"), text)
        elif stripped.startswith("- "):
            indent = len(line) - len(line.lstrip())  # 들여쓴 글머리표는 한 단계 안으로
            i, text = take_block(lines, i, stripped[2:])
            add_runs(doc.add_paragraph(style="List Bullet 2" if indent >= 2 else "List Bullet"), text)
        elif stripped.startswith("**Q."):
            # 질문은 답과 한 문단으로 붙지 않게 따로 낸다. 소리 내어 읽을 때
            # 질문에서 한 번 끊기는 게 이 문서의 용도에 맞는다.
            add_runs(doc.add_paragraph(), stripped)
        else:
            i, text = take_block(lines, i, stripped)
            add_runs(doc.add_paragraph(), text)
        i += 1

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run(f"출처: {src_label} (저장소가 원본이다)")
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY

    doc.save(out)
    print(f"저장: {out}")


def desktop_dir():
    """바탕화면의 개발일지 폴더. 사용자 이름을 박지 않는다(devlog_doc.py와 같은 이유)."""
    import glob

    for pattern in ("/mnt/c/Users/*/OneDrive/바탕 화면/개발일지", "/mnt/c/Users/*/Desktop/개발일지"):
        hits = sorted(glob.glob(pattern))
        if hits:
            return hits[0]
    raise SystemExit("바탕화면 개발일지 폴더를 못 찾았다 — 출력 경로를 직접 주세요.")


if __name__ == "__main__":
    args = sys.argv[1:]
    if len(args) < 2:
        raise SystemExit(__doc__.rsplit("사용:", 1)[-1].strip())
    if args[1] == "--desktop":
        out = f"{desktop_dir()}/{args[2]}"
    else:
        out = args[1]
    main(args[0], out)
